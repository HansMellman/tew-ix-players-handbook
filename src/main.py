import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from jinja2 import Environment, FileSystemLoader

# === Load and prepare the data ===
df = pd.read_excel("hidden\\tblHandbook.xlsx")
df = df.sort_values("OrderID")


def clean_text(text):
    if isinstance(text, str):
        return text.replace("_x000d_", "").replace("\r", "").replace("\n", "\n")
    return text


# === Word Document ===
word_doc = Document()

for _, row in df.iterrows():
    if row["SectionHeader"]:
        word_doc.add_heading(str(row["HeaderText"]), level=1)
    elif row["SectionID"] != 0:
        word_doc.add_heading(str(row["HeaderText"]), level=2)
    else:
        word_doc.add_heading(str(row["HeaderText"]), level=3)

    if pd.notnull(row["BodyText"]):
        para = word_doc.add_paragraph(clean_text(row["BodyText"]))
        para.style.font.size = Pt(11)

word_doc.save("handbook.docx")


# === PDF Document ===
class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, "TEW IX Handbook", ln=True, align="C")
        self.ln(10)

    def chapter_title(self, title, level):
        font_size = 14 - (level * 2)
        self.set_font("Arial", "B", font_size)
        self.multi_cell(0, 8, title)
        self.ln(2)

    def chapter_body(self, body):
        self.set_font("Arial", "", 11)
        self.multi_cell(0, 8, body)
        self.ln()


pdf = PDF()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.add_page()

for _, row in df.iterrows():
    level = 1 if row["SectionHeader"] else 2 if row["SectionID"] != 0 else 3
    pdf.chapter_title(clean_text(row["HeaderText"]), level)
    if pd.notnull(row["BodyText"]):
        pdf.chapter_body(clean_text(row["BodyText"]))

pdf.output("handbook.pdf")


def slugify(text):
    return re.sub(r"[^\w\-]+", "-", text.strip()).lower()


# html_template = """
# <!DOCTYPE html>
# <html lang="en">
# <head>
#   <meta charset="UTF-8">
#   <meta name="viewport" content="width=device-width, initial-scale=1.0">
#   <title>TEW IX Handbook</title>
#   <style>
#     body { font-family: Arial, sans-serif; padding: 2em; max-width: 900px; margin: auto; scroll-behavior: smooth; }
#     h1, h2, h3 { color: #2c3e50; }
#     h1 { border-bottom: 2px solid #ccc; }
#     .section { margin-bottom: 2em; }
#     .toc { background: #f9f9f9; border: 1px solid #ccc; padding: 1em; margin-bottom: 2em; }
#     .toc ul { list-style-type: none; padding-left: 1em; }
#     .toc li { margin: 0.3em 0; }

#     p {
#   word-wrap: break-word;
#   overflow-wrap: break-word;
#   line-height: 1.6;
#   }

#     #top-button {
#       display: none;
#       position: fixed;
#       bottom: 30px;
#       right: 30px;
#       z-index: 99;
#       background-color: #2c3e50;
#       color: white;
#       border: none;
#       padding: 12px 16px;
#       border-radius: 6px;
#       font-size: 14px;
#       cursor: pointer;
#       opacity: 0.8;
#     }
#     #top-button:hover {
#       background-color: #1a242f;
#     }
#   </style>
# </head>
# <body>
#   <a id="top"></a>
#   <h1>TEW IX Player's Handbook</h1>

#   <div class="toc">
#     <h2>Table of Contents</h2>
#     <ul>
#       {% for row in toc %}
#         <li style="margin-left: {{ (row.level - 1) * 20 }}px;">
#           <a href="#{{ row.id }}">{{ row.text }}</a>
#         </li>
#       {% endfor %}
#     </ul>
#   </div>

#   {% for row in rows %}
#     <div class="section">
#       <{{ row.heading_tag }} id="{{ row.id }}">{{ row.HeaderText }}</{{ row.heading_tag }}>
#       {% if row.BodyText %}
#       <p>{{ row.BodyText }}</p>
#       {% endif %}
#     </div>
#   {% endfor %}

#   <button onclick="scrollToTop()" id="top-button" title="Return to Top">↑ Top</button>

#   <script>
#     const topButton = document.getElementById("top-button");

#     window.onscroll = function() {
#       if (document.body.scrollTop > 400 || document.documentElement.scrollTop > 400) {
#         topButton.style.display = "block";
#       } else {
#         topButton.style.display = "none";
#       }
#     };

#     function scrollToTop() {
#       window.scrollTo({ top: 0, behavior: 'smooth' });
#     }
#   </script>
# </body>
# </html>
# """

html_template = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>TEW IX Handbook</title>
  <style>
    :root {
      --bg-color: #ffffff;
      --text-color: #2c3e50;
      --link-color: #2c3e50;
      --hover-color: #2980b9;
      --toc-bg: #f9f9f9;
      --toc-border: #ccc;
    }

    [data-theme="dark"] {
      --bg-color: #1e1e1e;
      --text-color: #f5f5f5;
      --link-color: #d1ecff;
      --hover-color: #74b9ff;
      --toc-bg: #2c2c2c;
      --toc-border: #555;
    }

    body {
      font-family: Arial, sans-serif;
      padding: 2em;
      max-width: 900px;
      margin: auto;
      background-color: var(--bg-color);
      color: var(--text-color);
      scroll-behavior: smooth;
    }

    h1, h2, h3 {
      color: var(--text-color);
    }

    h1 {
      border-bottom: 2px solid var(--toc-border);
    }

    .section {
      margin-bottom: 2em;
    }

    .toc {
      background: var(--toc-bg);
      border: 1px solid var(--toc-border);
      padding: 1em;
      margin-bottom: 2em;
    }

    .toc ul {
      list-style-type: none;
      padding-left: 1em;
    }

    .toc li {
      margin: 0.3em 0;
    }

    .toc a {
      text-decoration: underline;
      color: var(--link-color);
    }

    .toc a:hover {
      color: var(--hover-color);
    }

    p {
      word-wrap: break-word;
      overflow-wrap: break-word;
      line-height: 1.6;
    }

    #top-button {
      display: none;
      position: fixed;
      bottom: 30px;
      right: 30px;
      z-index: 99;
      background-color: var(--text-color);
      color: var(--bg-color);
      border: none;
      padding: 12px 16px;
      border-radius: 6px;
      font-size: 14px;
      cursor: pointer;
      opacity: 0.8;
    }

    #top-button:hover {
      background-color: var(--hover-color);
    }

    #searchBar {
      width: 60%;
      padding: 8px;
      margin-bottom: 1em;
    }
    
    #mode-toggle {
      padding: 10px 16px;
      font-size: 14px;
      border-radius: 6px;
      cursor: pointer;
      background-color: #2c3e50;
      color: white;
      border: none;
      margin-left: 10px;
    }
    #mode-toggle:hover {
      background-color: #1a242f;
    }

  </style>
</head>
<body data-theme="light">
  <a id="top"></a>
  <h1>TEW IX Player's Handbook</h1>

  <input type="text" id="searchBar" placeholder="Type to search...">
  <button id="mode-toggle" onclick="toggleTheme()">Dark/Light</button>

  <div class="toc">
    <h2>Table of Contents</h2>
    <ul id="toc-list">
      {% for row in toc %}
        <li style="margin-left: {{ (row.level - 1) * 20 }}px;">
          <a href="#{{ row.id }}">{{ row.text }}</a>
        </li>
      {% endfor %}
    </ul>
  </div>

  <div id="content">
    {% for row in rows %}
      <div class="section">
        <{{ row.heading_tag }} id="{{ row.id }}">{{ row.HeaderText }}</{{ row.heading_tag }}>
        {% if row.BodyText %}<p>{{ row.BodyText }}</p>{% endif %}
      </div>
    {% endfor %}
  </div>

  <button onclick="scrollToTop()" id="top-button" title="Return to Top">↑ Top</button>

  <script>
    const topButton = document.getElementById("top-button");
    const body = document.body;

    window.onscroll = function () {
      topButton.style.display = (document.documentElement.scrollTop > 400) ? "block" : "none";
    };

    function scrollToTop() {
      window.scrollTo({ top: 0, behavior: 'smooth' });
    }

    function toggleTheme() {
      const theme = body.getAttribute("data-theme");
      body.setAttribute("data-theme", theme === "dark" ? "light" : "dark");
    }

    // Basic search filter
    document.getElementById("searchBar").addEventListener("input", function (e) {
      const query = e.target.value.toLowerCase();
      const items = document.querySelectorAll("#toc-list li");
      items.forEach(li => {
        const text = li.textContent.toLowerCase();
        li.style.display = text.includes(query) ? "list-item" : "none";
      });
    });
  </script>
</body>
</html>
"""

env = Environment(loader=FileSystemLoader("."))
template = env.from_string(html_template)

rows = []
toc = []
for _, row in df.iterrows():
    tag = "h1" if row["SectionHeader"] else "h2" if row["SectionID"] != 0 else "h3"
    header_text = clean_text(row["HeaderText"])
    body_text = clean_text(row["BodyText"]) if pd.notnull(row["BodyText"]) else ""
    heading_id = slugify(header_text)

    rows.append(
        {
            "HeaderText": header_text,
            "BodyText": body_text,
            "heading_tag": tag,
            "id": heading_id,
        }
    )

    toc.append(
        {
            "id": heading_id,
            "text": header_text,
            "level": 1 if tag == "h1" else 2 if tag == "h2" else 3,
        }
    )

html_output = template.render(rows=rows, toc=toc)
Path("handbook.html").write_text(html_output, encoding="utf-8")
